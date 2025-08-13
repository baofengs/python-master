from docx import Document
import os

# 创建一个新的Word文档
doc = Document()

# 添加一些带拼音的文本
doc.add_paragraph("天地（tiān dì）玄黄（xuán huáng）")
doc.add_paragraph("宇宙（yǔ zhòu）洪荒（hóng huāng）")

# 确保docs目录存在
current_dir = os.path.dirname(os.path.abspath(__file__))
docs_dir = os.path.join(current_dir, "docs")
os.makedirs(docs_dir, exist_ok=True)

# 保存文档
input_docx = os.path.join(docs_dir, "input.docx")
doc.save(input_docx)
print(f"✅ 已创建测试文档: {input_docx}")
